import os
import json
from docx import Document # type: ignore

DOCS_FOLDER = "./documents"

def extract_text_from_docx(filepath):
  doc = Document(filepath)

  # Paragraphs ka text
  paragraphs = [para.text.strip() for para in doc.paragraphs]
  clean_paragraphs = [p for p in paragraphs if p]

  # Tables ka text
  table_texts = []
  for table in doc.tables:
    for row in table.rows:
      for cell in row.cells:
        if cell.text.strip():
          table_texts.append(cell.text.strip())

  all_text = clean_paragraphs + table_texts
  return " ".join(all_text)

def extrect_all_documents(folder):
  documents = {}

  for filename in sorted(os.listdir(folder)):
    if filename.endswith(".docx"):
      doc_id = filename.replace(".docx","")
      filepath = os.path.join(folder,filename)
      text = extract_text_from_docx(filepath)
      documents[doc_id] = text
      print(f"Extracted: {doc_id} ({len(text)} chars)")

  return documents

if __name__ == "__main__":
  print("Extracting text from all documents...\n")
  documents = extrect_all_documents(DOCS_FOLDER) 

  with open("documents.json", "w", encoding="utf-8") as f:
    json.dump(documents, f, indent=2, ensure_ascii=False)

  print(f"\n Done! {len(documents)} documents saved to documents.json")