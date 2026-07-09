from docx import Document
import os

def extract_text(file_path: str) -> str:
    doc = Document(file_path)

    paragraphs = [p.text.strip() for p in doc.paragraphs]
    clean_paragraphs = [p for p in paragraphs if p]

    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    table_texts.append(cell.text.strip())
    
    all_text = clean_paragraphs + table_texts
    return " ".join(all_text)